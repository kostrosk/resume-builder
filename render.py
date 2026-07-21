"""
render.py — turn selected content + style config into an ATS-safe .docx.

Supports inline markdown inside any bullet or text field:

    **bold**      bold
    *italic*      italic
    ***both***    bold italic
    __underline__ underline
    `code`        rendered in the body font, not a monospace substitute,
                  because ATS parsers mishandle font switches mid-run

Deliberately avoids: tables, text boxes, columns, headers/footers, images.
Every one of those breaks applicant tracking systems.
"""
import re, datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MONTHS = ["", "Jan", "Feb", "Mar", "Apr", "May", "Jun",
          "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]

# ***x*** | **x** | *x* | __x__ | `x`
INLINE = re.compile(r"(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*[^*\n]+?\*|__.+?__|`[^`\n]+?`)")


def fmt_date(v, style):
    """Accepts 2024-03, 'present', or free text. Free text passes through."""
    if v is None:
        return ""
    s = str(v).strip()
    if s.lower() in ("present", "current", "now"):
        return "Present"
    m = re.match(r"^(\d{4})-(\d{1,2})$", s)
    if m:
        y, mo = int(m.group(1)), int(m.group(2))
        if style.get("date_format", "MMM yyyy") == "yyyy":
            return str(y)
        return f"{MONTHS[mo]} {y}"
    m = re.match(r"^(\d{4})$", s)
    return m.group(1) if m else s


def date_range(job, style):
    a = fmt_date(job.get("start"), style)
    b = fmt_date(job.get("end"), style)
    sep = style.get("date_separator", " – ")
    return f"{a}{sep}{b}" if a and b else (a or b)


def _rule(par, color):
    """Horizontal line under a paragraph, via a bottom border."""
    p = par._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    bdr.append(bot)
    p.append(bdr)


class Renderer:
    def __init__(self, cfg):
        self.cfg = cfg
        self.s = cfg.get("style", {})
        self.size = self.s.get("size", {})
        self.color = self.s.get("color", {})
        self.space = self.s.get("spacing", {})
        self.doc = Document()
        base = self.doc.styles["Normal"]
        base.font.name = self.s.get("font", "Calibri")
        base.font.size = Pt(self.size.get("body", 10))
        # East-Asian font mapping, or Word silently substitutes
        rpr = base.element.get_or_add_rPr()
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts")
            rpr.append(rf)
        rf.set(qn("w:eastAsia"), self.s.get("font", "Calibri"))
        m = self.s.get("margins", {})
        for sec in self.doc.sections:
            sec.top_margin = Inches(m.get("top", 0.5))
            sec.bottom_margin = Inches(m.get("bottom", 0.5))
            sec.left_margin = Inches(m.get("left", 0.5))
            sec.right_margin = Inches(m.get("right", 0.5))

    # ---------------------------------------------------------------- text
    def _runs(self, par, text, size, color, bold=False, italic=False):
        for chunk in INLINE.split(text):
            if not chunk:
                continue
            b, i, u = bold, italic, False
            if chunk.startswith("***") and chunk.endswith("***"):
                chunk, b, i = chunk[3:-3], True, True
            elif chunk.startswith("**") and chunk.endswith("**"):
                chunk, b = chunk[2:-2], True
            elif chunk.startswith("__") and chunk.endswith("__"):
                chunk, u = chunk[2:-2], True
            elif chunk.startswith("*") and chunk.endswith("*") and len(chunk) > 2:
                chunk, i = chunk[1:-1], True
            elif chunk.startswith("`") and chunk.endswith("`"):
                chunk = chunk[1:-1]
            r = par.add_run(chunk)
            r.bold, r.italic, r.underline = b, i, u
            r.font.size = Pt(size)
            if color:
                r.font.color.rgb = RGBColor.from_string(color)

    def para(self, text, size=None, color=None, bold=False, italic=False,
             space_after=2, space_before=0, align=None, bullet=False, rule=False):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(space_before)
        pf.space_after = Pt(space_after)
        pf.line_spacing = self.s.get("line_spacing", 1.0)
        if bullet:
            ind = self.s.get("bullet_indent", 0.16)
            pf.left_indent = Inches(ind)
            pf.first_line_indent = Inches(-ind)
        if align:
            p.alignment = align
        size = size or self.size.get("body", 10)
        if bullet:
            ch = self.s.get("bullet_char", "•")
            r = p.add_run(f"{ch} ")
            r.font.size = Pt(size)
            if color:
                r.font.color.rgb = RGBColor.from_string(color)
        self._runs(p, text, size, color, bold, italic)
        if rule:
            _rule(p, self.color.get("section", "999999"))
        return p

    # ------------------------------------------------------------ sections
    def header(self, ident):
        self.para(ident.get("name", ""), size=self.size.get("name", 16),
                  color=self.color.get("name"), bold=True,
                  space_after=self.space.get("after_name", 1),
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        if ident.get("headline"):
            self.para(ident["headline"], size=self.size.get("headline", 9),
                      color=self.color.get("muted"), italic=True, space_after=1,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        bits = [ident.get(k) for k in ("location", "phone", "email")]
        # links: any number, any label. show: false benches one without
        # deleting it. Legacy flat keys still work.
        for lk in ident.get("links") or []:
            if lk.get("show", True) and (lk.get("url") or "").strip():
                bits.append(lk["url"].strip())
        for legacy in ("linkedin", "github"):
            if ident.get(legacy):
                bits.append(ident[legacy])
        sep = self.s.get("link_separator", "  |  ")
        line = sep.join([b for b in bits if b])
        self.para(line, size=self.size.get("contact", 9),
                  color=self.color.get("muted"),
                  space_after=self.space.get("after_contact", 6),
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    def section_title(self, title):
        if self.s.get("uppercase_sections", True):
            title = title.upper()
        self.para(title, size=self.size.get("section", 11),
                  color=self.color.get("section"), bold=True,
                  space_before=self.space.get("before_section", 6),
                  space_after=self.space.get("after_section", 3),
                  rule=self.s.get("section_rule", True))

    def job_block(self, job, bullets):
        name = job.get("company", "")
        if job.get("ticker"):
            name = f"{name} ({job['ticker']})"
        role = job.get("role", "")
        head = f"**{name}**" + (f" — {role}" if role else "")
        self.para(head, size=self.size.get("company", 10),
                  color=self.color.get("accent"),
                  space_after=self.space.get("after_company", 0))
        meta = date_range(job, self.s)
        if job.get("location"):
            meta = f"{meta}  |  {job['location']}"
        self.para(meta, size=self.size.get("dates", 9),
                  color=self.color.get("muted"), italic=True,
                  space_after=self.space.get("after_dates", 2))
        for b in bullets:
            self.para(b, size=self.size.get("body", 10),
                      color=self.color.get("body"), bullet=True,
                      space_after=self.space.get("after_bullet", 2))

    def tags(self, items):
        self.para("  ·  ".join(items), size=self.size.get("tags", 9),
                  color=self.color.get("body"), space_after=4)

    def text_block(self, body):
        for para in [p for p in body.split("\n") if p.strip()]:
            self.para(para.strip(), space_after=3)

    def save(self, path):
        self.doc.save(path)
        return path


def verify(path, expected):
    """Re-open and confirm every shipped line extracts. A .docx that does not
    re-parse is a resume no ATS can read."""
    d = Document(path)
    text = "\n".join(p.text for p in d.paragraphs)
    plain = [re.sub(r"[*_`]", "", e) for e in expected]
    missing = [e[:60] for e in plain if e[:45] not in text]
    return {
        "paragraphs": len(d.paragraphs),
        "chars": len(text),
        "tables": len(d.tables),
        "sections": len(d.sections),
        "missing": missing,
    }
