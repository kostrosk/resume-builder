"""
tailor.py — deterministic JD-tailored one-page resume.

  python tailor.py jd/crowdstrike.md

Reads ONLY `- [x]` bullets from master-resume.md. Unconfirmed content cannot
reach output; that is the enforcement point, not a convention.

Never reads source/draft/ — enforced below.

No model call unless --llm is passed (not implemented in v1).
"""
import os, re, sys, glob, argparse, datetime
from collections import defaultdict
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.abspath(__file__))
MASTER = os.path.join(ROOT, "master-resume.md")
SYN = os.path.join(ROOT, "config", "synonyms.yaml")
OUT = os.path.join(ROOT, "output")
LIN = os.path.join(ROOT, "lineage")
RES = os.path.join(ROOT, "source", "resumes")
FORBIDDEN = os.path.join(ROOT, "source", "draft")

STOP = set("""a an the and or of to in for with on at by from as is are was were be been being
this that these those it its our your their we i you they he she them us not no if then than
will would can could should may might must have has had do does did doing done more most other
such own same so too very just about into over under out up down off across per each any all
role position job candidate company team work working ability years year experience required
requirements responsibilities preferred qualifications plus strong excellent good ensure ensuring
including include includes etc via while within across also new using use used help helps""".split())


def die(m):
    print(f"ERROR: {m}", file=sys.stderr)
    sys.exit(1)


# --------------------------------------------------------------- stemming
def stem(w):
    for suf in ("ization", "izations", "ations", "ation", "ances", "ance", "ences",
                "ence", "ments", "ment", "ings", "ing", "ies", "ied", "ers", "er",
                "ed", "es", "s"):
        if len(w) > len(suf) + 3 and w.endswith(suf):
            return w[: -len(suf)]
    return w


def toks(text):
    ws = re.findall(r"[a-z0-9][a-z0-9+#/.\-]*", text.lower())
    return [stem(w) for w in ws if w not in STOP and len(w) > 2]


def phrases(text):
    """Bigrams + trigrams, stemmed, for multiword competency matching."""
    ws = [w for w in re.findall(r"[a-z0-9][a-z0-9+#/\-]*", text.lower()) if len(w) > 1]
    out = set()
    for n in (2, 3):
        for i in range(len(ws) - n + 1):
            g = ws[i:i + n]
            if g[0] in STOP or g[-1] in STOP:
                continue
            out.add(" ".join(stem(x) for x in g))
    return out


# --------------------------------------------------------------- synonyms
def load_syn():
    groups = []
    if not os.path.exists(SYN):
        return groups
    for ln in open(SYN, encoding="utf-8"):
        ln = ln.strip()
        if not ln or ln.startswith("#") or ":" not in ln:
            continue
        k, v = ln.split(":", 1)
        terms = [k.strip()] + [x.strip() for x in v.split(",") if x.strip()]
        groups.append({" ".join(stem(w) for w in t.lower().split()) for t in terms if t})
    return groups


SYNG = load_syn()


def expand(term_set):
    out = set(term_set)
    for g in SYNG:
        if term_set & g:
            out |= g
    return out


# ------------------------------------------------------------ master parse
def load_master():
    if not os.path.exists(MASTER):
        die("master-resume.md not found — run build_master.py first")
    roles, cur, section = [], None, ""
    confirmed = skipped = 0
    for ln in open(MASTER, encoding="utf-8"):
        t = ln.rstrip("\n")
        m = re.match(r"^##\s+(.+)", t)
        if m and not t.startswith("###"):
            section = m.group(1).strip()
            continue
        m = re.match(r"^###\s+(.+)", t)
        if m:
            head = m.group(1).strip()
            if "—" in head:
                emp, rest = head.split("—", 1)
                title, dates = (rest.split("|", 1) + [""])[:2]
            else:
                emp, title, dates = head, "", ""
            cur = {"employer": emp.strip(), "title": title.strip(),
                   "dates": dates.strip(), "section": section, "bullets": []}
            roles.append(cur)
            continue
        m = re.match(r"^\s*-\s*\[([ xX])\]\s*(.+)", t)
        if m and cur is not None:
            if m.group(1).lower() == "x":
                cur["bullets"].append(m.group(2).strip())
                confirmed += 1
            else:
                skipped += 1
    return [r for r in roles if r["bullets"]], confirmed, skipped


# ----------------------------------------------------------------- JD parse
SENIOR_SIGNALS = {
    "vp": 4, "vice president": 4, "svp": 4, "head of": 4, "chief": 4,
    "senior director": 4, "sr. director": 4, "executive": 3,
    "director": 3, "head": 3,
    "senior manager": 2, "sr. manager": 2, "manager": 2, "lead": 2,
    "principal": 2, "architect": 2,
    "analyst": 1, "engineer": 1, "specialist": 1, "associate": 1,
}
STRAT_WORDS = ("strategy", "operating model", "roadmap", "executive", "board",
               "council", "decision rights", "vision", "budget", "p&l",
               "organizational", "influence", "stakeholder", "transformation")
EXEC_WORDS = ("sql", "python", "dbt", "pipeline", "hands-on", "implement",
              "configure", "build", "develop", "automate", "script", "etl",
              "tagging", "masking", "query", "code")


def classify(jd_text, jd_title):
    low = (jd_title + " " + jd_text[:1500]).lower()
    best = 0
    for k, v in SENIOR_SIGNALS.items():
        if k in low:
            best = max(best, v)
    s = sum(jd_text.lower().count(w) for w in STRAT_WORDS)
    e = sum(jd_text.lower().count(w) for w in EXEC_WORDS)
    ratio = s / max(e, 1)
    if best >= 4 or (best == 3 and ratio > 1.6):
        prof = "DIRECTOR+/VP"
    elif best == 3 or (best == 2 and ratio > 2.0):
        prof = "DIRECTOR"
    elif best == 2:
        prof = "SENIOR MANAGER/LEAD"
    else:
        prof = "BELOW DIRECTOR"
    return prof, best, s, e, ratio


VERB_MAP = {
    "DIRECTOR+/VP": [("led", "directed"), ("built", "established"),
                     ("implemented", "institutionalized"), ("ran", "chaired"),
                     ("managed", "directed"), ("created", "established"),
                     ("designed", "established"), ("developed", "established")],
    "DIRECTOR": [("spearheaded", "led"), ("drove", "led"), ("created", "designed"),
                 ("engineered", "designed")],
    "SENIOR MANAGER/LEAD": [("directed", "built"), ("led", "built"),
                            ("established", "implemented"), ("chaired", "ran"),
                            ("institutionalized", "implemented"),
                            ("spearheaded", "built"), ("designed", "developed")],
    "BELOW DIRECTOR": [],
}


def apply_verbs(text, profile):
    """Swap only the leading verb. Facts are never touched."""
    for a, b in VERB_MAP.get(profile, []):
        m = re.match(rf"^\s*{a}\b", text, re.I)
        if m:
            rep = b.capitalize() if text[:1].isupper() else b
            return re.sub(rf"^\s*{a}\b", rep, text, count=1, flags=re.I), (a, b)
    return text, None


# ------------------------------------------------------------------ scoring
def jd_terms(jd_text):
    """Weight terms by prominence: requirement bullets and early text count more."""
    lines = jd_text.splitlines()
    w = defaultdict(float)
    for i, ln in enumerate(lines):
        pos = 1.6 if i < len(lines) * 0.35 else 1.0
        bul = 1.5 if re.match(r"^\s*[-*•]|\d+\.", ln) else 1.0
        for t in set(toks(ln)):
            w[t] += pos * bul
        for p in phrases(ln):
            w[p] += pos * bul * 2.0        # multiword competencies dominate
    return w


def score_bullet(text, jdw):
    bt = set(toks(text)) | phrases(text)
    bt = expand(bt)
    hit = {t: jdw[t] for t in bt if t in jdw}
    raw = sum(hit.values())
    if raw >= 14:
        s = 3
    elif raw >= 7:
        s = 2
    elif raw >= 2.5:
        s = 1
    else:
        s = 0
    return s, raw, sorted(hit, key=lambda x: -hit[x])[:6]


# ------------------------------------------------------------- page budget
CPL = 108          # chars per line at 10pt Calibri, 0.5in margins
MAX_LINES = 44     # body lines that fit one page after header/sections


def est_lines(s):
    return max(1, -(-len(s) // CPL))


def select(roles, jdw, profile):
    scored = []
    for ri, r in enumerate(roles):
        for b in r["bullets"]:
            s, raw, hits = score_bullet(b, jdw)
            scored.append({"role": ri, "text": b, "score": s, "raw": raw, "hits": hits})
    # recency weight: most recent role is the sales pitch
    n = len(roles)
    for x in scored:
        rec = 1.0 + 0.5 * ((n - x["role"]) / max(n, 1))
        x["rank"] = x["raw"] * rec
    scored.sort(key=lambda x: (-x["score"], -x["rank"]))

    used, chosen, per = 0, [], defaultdict(int)
    cap = {0: 12}                       # most relevant role gets the most air
    for x in scored:
        if x["score"] == 0:
            continue
        lim = cap.get(x["role"], 4 if x["role"] <= 2 else 2)
        if per[x["role"]] >= lim:
            continue
        need = est_lines(x["text"])
        if used + need > MAX_LINES:
            continue
        chosen.append(x)
        per[x["role"]] += 1
        used += need
    return chosen, scored, used


# ---------------------------------------------------------------- ATS score
def ats(chosen, jdw, jd_title, roles):
    top = sorted(jdw, key=lambda x: -jdw[x])[:60]
    body = " ".join(x["text"] for x in chosen) + " " + " ".join(
        f"{r['title']} {r['employer']}" for r in roles)
    have = expand(set(toks(body)) | phrases(body))
    got = [t for t in top if t in have]
    miss = [t for t in top if t not in have]
    cov = len(got) / max(len(top), 1)
    tmatch = len(set(toks(jd_title)) & have) / max(len(set(toks(jd_title))), 1)
    score = round(100 * (0.75 * cov + 0.25 * tmatch))
    return score, miss, got


# ------------------------------------------------------------------- docx
def write_docx(path, contact, summary, roles, chosen, profile):
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.5)
        s.left_margin = s.right_margin = Inches(0.6)

    def para(text, bold=False, size=10, space=2, align=None, bullet=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(space)
        if bullet:
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.first_line_indent = Inches(-0.18)
        if align:
            p.alignment = align
        r = p.add_run(("• " if bullet else "") + text)
        r.bold = bold
        r.font.size = Pt(size)
        return p

    para(contact["name"], bold=True, size=16, space=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(" | ".join(contact["line"]), size=9, space=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    if summary:
        para("SUMMARY", bold=True, size=11, space=3)
        para(summary, space=8)

    para("PROFESSIONAL EXPERIENCE", bold=True, size=11, space=3)
    by = defaultdict(list)
    for x in chosen:
        by[x["role"]].append(x)
    for ri, r in enumerate(roles):
        if ri not in by:
            continue
        para(f"{r['employer']} — {r['title']}", bold=True, space=0)
        para(r["dates"], size=9, space=2)
        for x in by[ri]:
            para(x["out"], bullet=True, space=2)

    doc.save(path)


def verify(path, chosen):
    """A .docx that does not re-parse fails ATS. Confirm every bullet extracts."""
    d = Document(path)
    text = "\n".join(p.text for p in d.paragraphs)
    missing = [x["out"][:60] for x in chosen if x["out"][:60] not in text]
    return len(d.paragraphs), len(text), missing, len(d.tables)


# -------------------------------------------------------------------- main
def _profile_name():
    """Name lives in gitignored config, not in published source."""
    p = os.path.join(ROOT, "config", "profile.yaml")
    if os.path.exists(p):
        for ln in open(p, encoding="utf-8"):
            if ln.strip().startswith("name:"):
                return ln.split(":", 1)[1].strip()
    return ""


def contact_info():
    email = phone = link = ""
    name = _profile_name() or "Your Name"
    loc = ""
    for fp in sorted(glob.glob(os.path.join(RES, "*.md")), key=os.path.getmtime, reverse=True):
        t = open(fp, encoding="utf-8").read()[:900]
        email = email or (re.search(r"[\w.\-]+@[\w\-]+\.\w+", t) or [""])[0] if re.search(r"[\w.\-]+@[\w\-]+\.\w+", t) else email
        m = re.search(r"\(?\d{3}\)?[ .\-]\d{3}[ .\-]\d{4}", t)
        phone = phone or (m.group(0) if m else "")
        m = re.search(r"(?:https?://)?(?:www\.)?linkedin\.com/in/[\w\-]+/?", t)
        link = link or (m.group(0) if m else "")
        m = re.search(r"^([A-Z][a-z]+,\s*[A-Z]{2})$", t, re.M)
        loc = loc or (m.group(1) if m else "")
        if email and phone and link:
            break
    return {"name": name, "line": [x for x in (loc, phone, email, link) if x]}


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("jd")
    ap.add_argument("--llm", action="store_true", help="not implemented in v1")
    a = ap.parse_args()

    if os.path.abspath(a.jd).startswith(os.path.abspath(FORBIDDEN)):
        die("build step must never read source/draft/")

    if not os.path.exists(a.jd):
        die(f"JD not found: {a.jd}")
    jd_text = open(a.jd, encoding="utf-8").read()
    name = os.path.splitext(os.path.basename(a.jd))[0]
    jd_title = next((l.strip("# ").strip() for l in jd_text.splitlines() if l.strip()), name)

    roles, confirmed, skipped = load_master()
    if not confirmed:
        die("no confirmed bullets — tick [x] boxes in master-resume.md first")

    jdw = jd_terms(jd_text)
    profile, sig_, s_, e_, ratio = classify(jd_text, jd_title)
    chosen, scored, used = select(roles, jdw, profile)

    swaps = []
    for x in chosen:
        x["out"], sw = apply_verbs(x["text"], profile)
        if sw:
            swaps.append((sw, x["text"][:50]))

    score, miss, got = ats(chosen, jdw, jd_title, roles)

    os.makedirs(OUT, exist_ok=True)
    os.makedirs(LIN, exist_ok=True)
    dp = os.path.join(OUT, f"{name}.docx")
    write_docx(dp, contact_info(), "", roles, chosen, profile)
    nparas, nchars, missing, ntables = verify(dp, chosen)

    with open(os.path.join(OUT, f"{name}-gaps.md"), "w", encoding="utf-8") as f:
        f.write(f"# Gaps — {jd_title}\n\n")
        f.write(f"- profile detected: **{profile}** (seniority signal {sig_}, "
                f"strategy/execution {s_}/{e_} = {ratio:.2f})\n")
        f.write(f"- ATS score: **{score}/100**\n")
        f.write(f"- bullets: {len(chosen)} shipped of {confirmed} confirmed "
                f"({skipped} unconfirmed excluded by gate)\n\n")
        if profile == "BELOW DIRECTOR":
            f.write("> **Seniority flag:** this JD reads below Director. "
                    "GOAL.md says flag rather than write down to it.\n\n")
        if score < 70:
            f.write("## Below 70 — real coverage gaps\n\n")
            f.write("No eligible confirmed bullet matches these. This is an "
                    "experience gap or a confirmation gap, not a wording problem.\n\n")
        f.write("## JD terms with no matching bullet\n\n")
        for t in miss[:40]:
            f.write(f"- {t}  (weight {jdw[t]:.1f})\n")
        f.write("\n## Covered\n\n")
        f.write(", ".join(got[:40]) + "\n")
        if swaps:
            f.write("\n## Verb framing applied (facts unchanged)\n\n")
            for (a_, b_), snip in swaps:
                f.write(f"- `{a_}` -> `{b_}` — {snip}…\n")

    with open(os.path.join(LIN, f"{name}.md"), "w", encoding="utf-8") as f:
        f.write(f"# Lineage — {name}\n\n| # | Bullet | Role | Score | JD terms hit |\n|---|---|---|---|---|\n")
        for i, x in enumerate(chosen, 1):
            r = roles[x["role"]]
            f.write(f"| {i} | {x['out'][:80].replace('|','\\|')}… | {r['employer']} "
                    f"{r['dates']} | {x['score']} | {', '.join(x['hits'])} |\n")

    print(f"profile   {profile}   (seniority {sig_}, strat/exec {s_}/{e_})")
    print(f"bullets   {len(chosen)} shipped / {confirmed} confirmed / {skipped} gated out")
    print(f"ATS       {score}/100" + ("   <-- below 70" if score < 70 else ""))
    print(f"lines     {used}/{MAX_LINES}")
    print(f"docx      {nparas} paragraphs, {nchars} chars, {ntables} tables, "
          f"{'ALL TEXT EXTRACTS' if not missing else f'{len(missing)} MISSING'}")
    print(f"\n  {dp}\n  {os.path.join(OUT, name + '-gaps.md')}\n  {os.path.join(LIN, name + '.md')}")
    if miss[:8]:
        print("\ntop uncovered JD terms: " + ", ".join(miss[:8]))


if __name__ == "__main__":
    main()
